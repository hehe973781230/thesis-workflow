#!/usr/bin/env python3
"""
MBA论文 Word 生成脚本 - 严格按 MBA 格式规范（方案A：格式写死）

规范来源：mba-thesis-workflow/SKILL.md Phase 3 标准 + 用户补充规范

【格式常量】
中文摘要：标题"摘 要"，黑体16磅加粗居中，内容小四12磅宋体，行距20磅
英文摘要：Abstract，Arial 16磅加粗居中，内容TNR 12磅，行距20磅
目录：跳过
各章标题："第1章  绪论"（空两格），黑体16磅，段前24磅段后18磅，每章另起一页
一级节标题："1.2  ×××"，黑体14磅，段前24磅段后6磅
二级节标题："1.2.1  ×××"，黑体13磅，段前12磅段后6磅
三级节标题："(1) ×××"，宋体12磅，与正文同段
正文：宋体12磅，两端对齐，首行缩进2字符，段前段后0磅，行距20磅
三线表：顶线1.5磅/表头底线0.75磅/底线0.5磅，无竖线
参考文献：中英文分编排序，中文在前，[1]序号
附录/致谢：标题同章标题
页码：封面无页码 → 摘要/目录罗马数字 → 正文阿拉伯数字续前编号
"""

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, sys, os

# ============ MBA 格式常量 ============
FONT_BODY_CN    = '宋体'
FONT_HEADING    = '黑体'
FONT_ENGLISH   = 'Times New Roman'
FONT_ENGLISH_AB= 'Arial'

SZ_BODY   = Pt(12)    # 小四=12磅
SZ_H1     = Pt(16)    # 一级标题16磅
SZ_H2     = Pt(14)    # 二级节标题14磅
SZ_H3     = Pt(13)    # 三级节标题13磅
SZ_H4     = Pt(12)    # 四级节标题12磅
SZ_SMALL  = Pt(10.5)  # 图注/表注10.5磅
SZ_REF    = Pt(12)    # 参考文献12磅

LINE_20   = Pt(20)    # 正文行距20磅
LINE_SGL  = Pt(15.6)  # 单倍行距（小四12磅≈15.6磅）

S_BEFORE_H1 = Pt(24); S_AFTER_H1 = Pt(18)
S_BEFORE_H2 = Pt(24); S_AFTER_H2 = Pt(6)
S_BEFORE_H3 = Pt(12); S_AFTER_H3 = Pt(6)
S_BEFORE_H4 = Pt(0);  S_AFTER_H4 = Pt(0)

FIRST_INDENT = Pt(42)  # 首行缩进2个汉字符（全角约21pt×2≈42pt）

BORDER_TOP    = '24'   # 顶线1.5磅
BORDER_HEADER = '12'   # 表头底线0.75磅
BORDER_BOTTOM = '8'    # 底线0.5磅

BLACK = RGBColor(0x00, 0x00, 0x00)

# ============ 工具函数 ============

def set_run(run, size, fname, bold=False, color=BLACK):
    run.font.size = size
    run.font.name = fname
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fname)
    rFonts.set(qn('w:ascii'), fname)
    rFonts.set(qn('w:hAnsi'), fname)
    rPr.append(rFonts)

def set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             line=None, sb=None, sa=None, fi=None, li=None):
    pf = para.paragraph_format
    pf.alignment = align
    if line is not None: pf.line_spacing = line
    if sb is not None:   pf.space_before = sb
    if sa is not None:   pf.space_after = sa
    if fi is not None:   pf.first_line_indent = fi
    if li is not None:   pf.left_indent = li

def blank(doc):
    p = doc.add_paragraph()
    set_para(p, line=Pt(6), sb=Pt(0), sa=Pt(0))
    return p

def strip_bold(t):
    return re.sub(r'\*\*([^*]+)\*\*', r'\1', t)

def is_tbl_sep(line):
    cols = [c.strip() for c in line.split('|')[1:-1]]
    return bool(cols) and all(re.match(r'^:?-+:?$', c) for c in cols)

def set_three_line(table):
    """三线表：顶线+底线，内部无"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    b = OxmlElement('w:tblBorders')

    top = OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),BORDER_TOP); top.set(qn('w:color'),'000000')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),BORDER_BOTTOM); bot.set(qn('w:color'),'000000')
    ih = OxmlElement('w:insideH'); ih.set(qn('w:val'),'none')
    iv = OxmlElement('w:insideV'); iv.set(qn('w:val'),'none')

    b.append(top); b.append(bot); b.append(ih); b.append(iv)
    for e in tblPr.findall(qn('w:tblBorders')): tblPr.remove(e)
    tblPr.append(b)

def set_header_bot_border(table):
    """表头行底部0.75磅"""
    if not table.rows: return
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),BORDER_HEADER); bot.set(qn('w:color'),'000000')
        for edge in ('top','left','right'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'),'none'); tcB.append(e)
        tcB.append(bot)
        for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
        tcPr.append(tcB)

def apply_tbl_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_three_line(table)
    if table.rows:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run(run, SZ_SMALL, FONT_BODY_CN)
        set_header_bot_border(table)

# ============ 目录插入函数 ============

def _insert_toc(doc):
    """
    在英文摘要后插入目录页。
    插入"目 录"标题 + TOC 域代码，Word 打开时会提示更新域生成目录。
    """
    # 目录页分页符
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    set_para(p_break, sb=Pt(0), sa=Pt(0))

    # "目 录" 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p_title, line=LINE_SGL, sb=S_BEFORE_H1, sa=S_AFTER_H1)
    run = p_title.add_run('目 录')
    set_run(run, SZ_H1, FONT_HEADING, bold=True)

    # TOC 域代码：基于 Heading 1-3 自动生成目录
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r' TOC \o "1-3" \h \z \u '
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    toc_run._r.append(fldChar_begin)
    toc_run._r.append(instrText)
    toc_run._r.append(fldChar_end)
    set_para(toc_para, line=LINE_20, sb=Pt(0), sa=Pt(0))


def _copy_proposal_cover(doc, proposal_path):
    """
    从开题报告 .docx 复制封面内容到论文 docx 第一页。
    复制范围：第一个非空段落开始，到"研究背景与研究问题"之前结束。
    完整保留原段落格式：字体/字号/加粗/颜色/对齐/缩进/段间距。
    """
    from docx import Document as DocxDoc
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    prop_doc = DocxDoc(proposal_path)
    paras = prop_doc.paragraphs

    # 找到封面内容起始位置（第一个非空段落）和结束位置（"研究背景"段落前）
    start_idx = None
    end_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t and start_idx is None:
            start_idx = i  # 第一个非空段落
        if ('研究背景与研究问题' in t or
                ('研究背景' in t and '问题' in t) or
                t.startswith('1.') or t.startswith('一、')):
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print('⚠️ 未找到封面起止位置，跳过封面复制')
        return

    for i in range(start_idx, end_idx):
        p = paras[i]
        t = p.text
        new_p = doc.add_paragraph()
        # 复制段落对齐
        new_p.alignment = p.alignment if p.alignment else WD_ALIGN_PARAGRAPH.LEFT
        # 复制段落格式（缩进、段间距）
        pf_src = p.paragraph_format
        pf_dst = new_p.paragraph_format
        if pf_src.line_spacing:
            pf_dst.line_spacing = pf_src.line_spacing
        if pf_src.space_before is not None:
            pf_dst.space_before = pf_src.space_before
        if pf_src.space_after is not None:
            pf_dst.space_after = pf_src.space_after
        if pf_src.first_line_indent is not None:
            pf_dst.first_line_indent = pf_src.first_line_indent

        if not t.strip():
            continue

        # 复制文字块（逐 run 保留完整格式）
        for run in p.runs:
            new_run = new_p.add_run(run.text)
            if run.font.name:
                new_run.font.name = run.font.name
                rPr = new_run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), run.font.name)
                rFonts.set(qn('w:ascii'), run.font.name)
                rFonts.set(qn('w:hAnsi'), run.font.name)
                rPr.append(rFonts)
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.bold is not None:
                new_run.font.bold = run.font.bold
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb


# ============ 入口校验 ============

def preflight(md_path):
    issues = []
    md_dir = os.path.dirname(md_path) or '.'
    bn = os.path.basename(md_path)

    possible = [os.path.join(md_dir, bn.replace('.md', t)) for t in ['_审核.md','_审核O.md','_审核H.md']]
    rf = None
    for p in possible:
        if os.path.exists(p): rf = p; break
    if not rf:
        pre = re.match(r'(论文[^\.]+)', bn)
        pp = pre.group(1) if pre else ''
        cands = [os.path.join(md_dir,f) for f in os.listdir(md_dir)
                 if f.endswith('.md') and '_审核' in f and pp in f]
        # Level 2: fuzzy match - extract short keyword (first 2-4 Chinese chars)
        if not cands and pre:
            short_kw = re.search(r'论文_([^_]+)', bn)
            if short_kw:
                kw = short_kw.group(1)
                cands = [os.path.join(md_dir,f) for f in os.listdir(md_dir)
                         if f.endswith('.md') and '_审核' in f and kw in f]
        if cands:
            cands.sort(key=lambda x: os.path.getmtime(x), reverse=True)
            rf = cands[0]

    if not rf:
        issues.append("❌ 未找到审核报告，请先完成 Review Agent 终审")
        return False, issues

    with open(rf, 'r', encoding='utf-8') as f:
        rc = f.read()
    reds = re.findall(r'🔴\s*项统计[：:]\s*(\d+)', rc)
    if reds and int(reds[0]) > 0:
        issues.append(f"❌ 审核报告存在{reds[0]}条🔴问题")
    if '✅ 通过' not in rc and '综合评级' in rc:
        issues.append("❌ 审核未通过")
    if issues:
        return False, issues

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    ch = len(re.findall(r'^#\s+第.+章', content, re.MULTILINE))
    bb = len(re.findall(r'^(?!#).+\*\*[^*]+\*\*', content, re.MULTILINE))
    wc = len(content)
    print(f"✅ 通过 | 审核：{os.path.basename(rf)} | 章节：{ch} | 字数：{wc}")
    if bb: print(f"   加粗残留：{bb}处（自动清除）")
    return True, issues

# ============ Word 生成 ============

def md_to_docx(md_path, docx_path, proposal_docx_path=None):
    doc = docx.Document()
    ns = doc.styles['Normal']
    ns.font.name = FONT_BODY_CN
    ns.font.size = SZ_BODY
    ns.paragraph_format.line_spacing = LINE_20

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    tbl_rows = []
    ref_cn = []; ref_en = []
    in_ref = False

    while i < len(lines):
        line = lines[i].rstrip('\n\r')
        i += 1

        if line.strip() == '===END===': continue
        if is_tbl_sep(line): continue

        if line.startswith('|'):
            cols = [strip_bold(c.strip()) for c in line.split('|')[1:-1]]
            tbl_rows.append(cols)
            continue

        if tbl_rows:
            _flush_tbl(doc, tbl_rows); tbl_rows = []

        if not line.strip(): continue

        # ---- 特殊章节检测 ----
        if re.match(r'^#{1,2}\s*摘要\s*$', line):
            _flush_tbl(doc, tbl_rows) if tbl_rows else None; tbl_rows = []
            i = _abs_cn(doc, lines, i)
            continue

        if re.match(r'^#{1,2}\s*英文摘要', line):
            _flush_tbl(doc, tbl_rows) if tbl_rows else None; tbl_rows = []
            i = _abs_en(doc, lines, i)
            continue

        if re.match(r'^#{1,2}\s*目录\s*$', line):
            while i < len(lines) and not re.match(r'^#\s+第.+章', lines[i]): i += 1
            continue

        if re.match(r'^#{1,2}\s*致谢\s*$', line):
            _flush_tbl(doc, tbl_rows) if tbl_rows else None; tbl_rows = []
            _chapter_title(doc, '致谢')
            i = _body_until_next(lines, i)
            continue

        am = re.match(r'^#{1,2}\s*附录([A-Z])?\s*$', line)
        if am:
            _flush_tbl(doc, tbl_rows) if tbl_rows else None; tbl_rows = []
            _appendix_title(doc, am.group(1) or '')
            i = _body_until_next(lines, i)
            continue

        # 各章标题（兼容 ## 第X章 和 # 第X章 两种格式）
        cm2 = re.match(r'^##\s+(第[一二三四五六七八九十\d]+章)\s+(.+)$', line)
        if cm2:
            raw = f'{cm2.group(1)}  {cm2.group(2)}'
            _chapter_title(doc, raw); continue

        cm = re.match(r'^#\s+(第[一二三四五六七八九十\d]+章)\s+(.+)$', line)
        if cm:
            raw = f'{cm.group(1)}  {cm.group(2)}'
            _chapter_title(doc, raw); continue

        # 一级节标题
        s1 = re.match(r'^##\s+(\d+\.\d+)\s+(.+)$', line)
        if s1:
            _sec1(doc, s1.group(1), s1.group(2)); continue

        # 二级节标题
        s2 = re.match(r'^###\s+(\d+\.\d+\.\d+)\s+(.+)$', line)
        if s2:
            _sec2(doc, s2.group(1), s2.group(2)); continue

        # 三级节标题
        s3 = re.match(r'^####\s+(\([^)]+\))\s*(.+)$', line)
        if s3:
            _sec3(doc, s3.group(1), s3.group(2)); continue

        # 参考文献
        if re.match(r'^#{1,2}\s+参考文献', line):
            _flush_tbl(doc, tbl_rows) if tbl_rows else None; tbl_rows = []
            in_ref = True; continue

        if in_ref and line.strip() and not line.startswith('#'):
            is_cn = bool(re.search(r'[\u4e00-\u9fff]', line))
            rt = re.sub(r'^\[\d+\]\s*', '', strip_bold(line.strip()))
            if is_cn: ref_cn.append(rt)
            else: ref_en.append(rt)
            continue
        elif in_ref and (not line.strip() or line.startswith('#')):
            _flush_refs(doc, ref_cn, ref_en)
            ref_cn = []; ref_en = []; in_ref = False

        # 正文段落
        para = doc.add_paragraph()
        set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
                 sb=Pt(0), sa=Pt(0), fi=FIRST_INDENT)
        run = para.add_run(strip_bold(line))
        set_run(run, SZ_BODY, FONT_BODY_CN)

    if tbl_rows: _flush_tbl(doc, tbl_rows)
    if ref_cn or ref_en: _flush_refs(doc, ref_cn, ref_en)

    # ============ 格式校验（规则型问题） ============
    warnings = validate_md_format(lines)
    if warnings:
        print('\n⚠️  Markdown 格式校验警告：')
        for w in warnings:
            print(f'  - {w}')

    doc.save(docx_path)
    return True

# ============ 格式校验函数（规则型问题，脚本自动检测） ============

def validate_md_format(lines):
    """
    规则型格式校验，发现问题打印警告但不中断生成。
    返回警告列表（用于汇总输出）。
    
    规则来源：MBA论文规范（用户补充版）
    """
    warnings = []
    prev_line = ''
    prev_prev_line = ''
    in_ref = False

    for i, line in enumerate(lines):
        ln = i + 1
        line_stripped = line.rstrip('\n\r')

        # ============ 格式类规则 ============
        # 规则1：禁止 ## 第X章 / ### 第X.X 混合格式
        if re.match(r'^##\s+第[一二三四五六七八九十\d]+章', line_stripped):
            warnings.append(
                f'第{ln}行：发现「## 第X章」格式，章标题应使用 # 第X章（行内容：{line_stripped[:30]}）'
            )
        if re.match(r'^###\s+第\d+\.\d+', line_stripped):
            warnings.append(
                f'第{ln}行：发现「### 第X.X」格式，节标题应使用 ## X.X（行内容：{line_stripped[:30]}）'
            )
        if re.match(r'^####\s+第', line_stripped):
            warnings.append(
                f'第{ln}行：发现「#### 第」格式，小节标题应使用 ### X.X.X（行内容：{line_stripped[:30]}）'
            )

        # 规则2：表格标题应在表格上方（检测表格前一行是否为标题格式）
        if line_stripped.startswith('|') and prev_line.strip():
            if not is_tbl_sep(prev_line) and not prev_line.startswith('|'):
                if not re.match(r'^#{1,3}\s+', prev_line):
                    warnings.append(
                        f'第{ln}行：表格前一行（第{ln-1}行）可能为标题但未使用「#」或「##」标记'
                    )

        # ============ 引用格式规则（规范5.3） ============
        # 检测上标式引用 [数字] 在正文中出现（规范5.3明确禁止）
        if re.match(r'^#{1,2}\s*参考文献', line_stripped):
            in_ref = True
        if in_ref and not line_stripped.startswith('#') and line_stripped.strip():
            pass  # 参考文献列表内容不检查
        if in_ref and (not line_stripped.strip() or line_stripped.startswith('#')):
            in_ref = False
        # 正文（非参考文献列表）中出现 [数字] 上标引用格式
        if not in_ref and not line_stripped.startswith('#'):
            sup_refs = re.findall(r'(?<![\w\[\]])\[(\d+)\]', line_stripped)
            if sup_refs:
                nums = ', '.join(sup_refs)
                warnings.append(
                    f'第{ln}行：正文出现上标编号式引用[{nums}]，规范5.3要求作者年制引用（如王明，2018）'
                )

        prev_prev_line = prev_line
        prev_line = line_stripped

    return warnings


# ============ 子函数 ============

def _flush_tbl(doc, rows):
    if not rows: return
    rn = len(rows); cn = max(len(r) for r in rows) if rows else 0
    tbl = doc.add_table(rows=rn, cols=cn)
    apply_tbl_style(tbl)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            if ci < cn:
                cell = tbl.rows[ri].cells[ci]
                cell.text = ct
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        set_run(run, SZ_SMALL, FONT_BODY_CN)

def _abs_cn(doc, lines, start):
    """中文摘要"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H1, sa=S_AFTER_H1)
    run = p.add_run()
    run.text = '摘 要'
    set_run(run, SZ_H1, FONT_HEADING, bold=True)

    content = []
    j = start
    while j < len(lines):
        line = lines[j].rstrip('\n\r')
        if re.match(r'^#{1,2}\s*英文摘要', line): break
        if line.strip() and not line.startswith('#'):
            content.append(strip_bold(line.strip()))
        j += 1

    for text in content:
        para = doc.add_paragraph()
        set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
                 sb=Pt(0), sa=Pt(0), fi=FIRST_INDENT)
        run = para.add_run(text)
        set_run(run, SZ_BODY, FONT_BODY_CN)

    kw = doc.add_paragraph()
    set_para(kw, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
             sb=Pt(0), sa=Pt(0), fi=FIRST_INDENT)
    run = kw.add_run('关键词 ')
    set_run(run, SZ_BODY, FONT_BODY_CN, bold=True)
    for text in content:
        m = re.search(r'关键词[：:]\s*(.+)', text)
        if m:
            run2 = kw.add_run(m.group(1))
            set_run(run2, SZ_BODY, FONT_BODY_CN)
            break
    return j

def _abs_en(doc, lines, start):
    """英文摘要"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H1, sa=S_AFTER_H1)
    run = p.add_run()
    run.text = 'Abstract'
    set_run(run, SZ_H1, FONT_ENGLISH_AB, bold=True)

    content = []
    j = start
    while j < len(lines):
        line = lines[j].rstrip('\n\r')
        if re.match(r'^#{1,2}\s*目录', line): break
        if line.strip() and not line.startswith('#'):
            content.append(strip_bold(line.strip()))
        j += 1

    for text in content:
        para = doc.add_paragraph()
        set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
                 sb=Pt(0), sa=Pt(0), fi=FIRST_INDENT)
        run = para.add_run(text)
        set_run(run, SZ_BODY, FONT_ENGLISH)

    kw = doc.add_paragraph()
    set_para(kw, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
             sb=Pt(0), sa=Pt(0), fi=FIRST_INDENT)
    run = kw.add_run('Key Words ')
    set_run(run, SZ_BODY, FONT_ENGLISH, bold=True)
    for text in content:
        m = re.search(r'[Kk]ey\s*[Ww]ords[：:]\s*(.+)', text)
        if m:
            run2 = kw.add_run(m.group(1))
            set_run(run2, SZ_BODY, FONT_ENGLISH)
            break
    return j

def _chapter_title(doc, raw_text):
    """各章标题：黑体16磅+段前24磅段后18磅，标题后分页（避免第一页空白）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H1, sa=S_AFTER_H1)
    run = p.add_run()
    run.text = raw_text
    set_run(run, SZ_H1, FONT_HEADING, bold=True)
    # 标题后分页，避免第一页空白
    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

def _appendix_title(doc, letter):
    """附录标题：标题后分页"""
    title = f'附录{letter}' if letter else '附录'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H1, sa=S_AFTER_H1)
    run = p.add_run()
    run.text = title
    set_run(run, SZ_H1, FONT_HEADING, bold=True)
    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

def _sec1(doc, idx, title):
    """一级节标题：1.2  ×××，黑体14磅，段前24磅段后6磅"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H2, sa=S_AFTER_H2)
    run = p.add_run()
    run.text = f'{idx}  {title}'
    set_run(run, SZ_H2, FONT_HEADING, bold=True)

def _sec2(doc, idx, title):
    """二级节标题：1.2.1  ×××，黑体13磅，段前12磅段后6磅"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, line=LINE_SGL, sb=S_BEFORE_H3, sa=S_AFTER_H3)
    run = p.add_run()
    run.text = f'{idx}  {title}'
    set_run(run, SZ_H3, FONT_HEADING, bold=True)

def _sec3(doc, idx, title):
    """三级节标题：(1) ×××，宋体12磅，与正文同段"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, line=LINE_20, sb=Pt(0), sa=Pt(0))
    run = p.add_run()
    run.text = f'{idx} {title}'
    set_run(run, SZ_H4, FONT_BODY_CN)

def _body_until_next(lines, start):
    """附录/致谢的正文，直到下一章或参考文献"""
    j = start
    while j < len(lines):
        line = lines[j].rstrip('\n\r')
        j += 1
        if re.match(r'^#\s+第.+章', line) or re.match(r'^#{1,2}\s*参考文献', line):
            return j - 1
        if not line.strip() or line.startswith('#'): continue
        if is_tbl_sep(line): continue
    return j

def _flush_refs(doc, cn, en):
    """参考文献输出"""
    if cn:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, line=LINE_SGL, sb=Pt(12), sa=Pt(6))
        run = p.add_run()
        run.text = '中文参考文献'
        set_run(run, SZ_H2, FONT_HEADING, bold=True)
        for ref in cn:
            para = doc.add_paragraph()
            set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
                     sb=Pt(0), sa=Pt(0), fi=Pt(0))
            run = para.add_run(ref)
            set_run(run, SZ_REF, FONT_BODY_CN)
    if en:
        if cn: blank(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, line=LINE_SGL, sb=Pt(12), sa=Pt(6))
        run = p.add_run()
        run.text = '英文参考文献'
        set_run(run, SZ_H2, FONT_HEADING, bold=True)
        for ref in en:
            para = doc.add_paragraph()
            set_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_20,
                     sb=Pt(0), sa=Pt(0), fi=Pt(0))
            run = para.add_run(ref)
            set_run(run, SZ_REF, FONT_ENGLISH)

# ============ 主入口 ============

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python3 md2docx_strict.py <论文.md> <输出.docx> [开题报告.docx]")
        sys.exit(1)
    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    proposal_path = sys.argv[3] if len(sys.argv) > 3 else None
    if not os.path.exists(md_path):
        print(f"❌ 文件不存在: {md_path}"); sys.exit(1)
    print("=== MBA Word 生成（Phase 3 规范）===")
    ok, issues = preflight(md_path)
    if not ok:
        for iss in issues: print(iss)
        sys.exit(1)
    if md_to_docx(md_path, docx_path, proposal_path):
        print(f"✅ 已生成: {docx_path}")
    else:
        print("❌ 生成失败"); sys.exit(1)